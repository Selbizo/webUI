import os
import re
import json
import uuid
import mimetypes
import logging
from pathlib import Path
from openai import OpenAI
from django.conf import settings as django_settings
from django.contrib.auth.decorators import login_required

BASE_DIR = Path(django_settings.BASE_DIR)
CONFIG_FILE = BASE_DIR / 'workspace_config.json'

def get_workspace_path():
    try:
        if CONFIG_FILE.exists():
            with open(CONFIG_FILE, 'r') as f:
                config = json.load(f)
                path = config.get('workspace_path', str(BASE_DIR))
                if os.path.isdir(path):
                    return path
    except Exception:
        pass
    return str(BASE_DIR)

def save_workspace_path(path):
    path = os.path.abspath(path)
    if not os.path.isdir(path):
        return False, 'Директория не существует'
    if not os.access(path, os.R_OK | os.W_OK):
        return False, 'Нет прав доступа'
    try:
        config = {'workspace_path': path}
        with open(CONFIG_FILE, 'w') as f:
            json.dump(config, f)
        return True, path
    except Exception as e:
        return False, str(e)

PROJECT_ROOT = get_workspace_path()
from django.http import JsonResponse, HttpResponseForbidden
from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

LM_STUDIO_BASE_URL = os.environ.get('LM_STUDIO_BASE_URL', 'http://localhost:1234/v1')
DEFAULT_MODEL = os.environ.get('DEFAULT_MODEL', '')

client = OpenAI(
    base_url=LM_STUDIO_BASE_URL,
    api_key='not-needed'
)

def get_available_models():
    try:
        models = client.models.list()
        model_ids = [model.id for model in models.data]
        print(f"Found models: {model_ids}")
        return model_ids
    except Exception as e:
        print(f"Error getting models: {e}")
        return []

def index(request):
    if not request.user.is_authenticated:
        return render(request, 'chat/login.html')
    return render(request, 'chat/index.html')

@csrf_exempt
def login_view(request):
    if request.method == 'GET':
        return render(request, 'chat/login.html')
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            username = data.get('username')
            password = data.get('password')
            
            print(f"Login attempt for user: {username}")
            
            user = authenticate(request, username=username, password=password)
            
            if user is not None:
                print(f"User {username} authenticated successfully")
                login(request, user)
                return JsonResponse({'success': True})
            else:
                print(f"Authentication failed for user: {username}")
                return JsonResponse({'success': False, 'error': 'Неверные учетные данные'}, status=401)
        except Exception as e:
            print(f"Login error: {e}")
            return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
def logout_view(request):
    logout(request)
    return redirect('/')

def health_check(request):
    models = get_available_models()
    return JsonResponse({
        'status': 'ready',
        'lm_studio_url': LM_STUDIO_BASE_URL,
        'available_models': models,
        'user': request.user.username if request.user.is_authenticated else None,
        'model_count': len(models)
    })

@csrf_exempt
@require_http_methods(['POST'])
def chat_api(request):
    try:
        data = json.loads(request.body)
        messages = data.get('messages', [])
        model = data.get('model', DEFAULT_MODEL) or (get_available_models()[0] if get_available_models() else '')
        
        if not model:
            return JsonResponse({'error': 'No model specified and no models available'}, status=500)
        
        session_id = data.get('session_id')
        context_files = data.get('context_files', [])
        
        if session_id and request.user.is_authenticated:
            from .models import ChatSession
            try:
                session = ChatSession.objects.get(session_id=session_id, user=request.user)
                stored_messages = session.messages
                messages = stored_messages + messages
            except ChatSession.DoesNotExist:
                pass
        
        if context_files:
            context_text = '\n\n'.join(
                f'=== Файл: {cf["path"]} ===\n{cf["content"]}'
                for cf in context_files
            )
            system_msg = f'Вот контекстные файлы для анализа:\n\n{context_text}\n\nПроанализируй эти файлы в контексте диалога.'
            messages = [{'role': 'system', 'content': system_msg}] + messages
        
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=data.get('temperature', 0.7),
            max_tokens=data.get('max_tokens', -1),
            stream=data.get('stream', False)
        )
        
        if data.get('stream', False):
            from django.http import StreamingHttpResponse
            import threading
            import queue

            result_queue = queue.Queue()
            
            def read_stream():
                try:
                    full_content = ''
                    for chunk in response:
                        if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                            content_chunk = chunk.choices[0].delta.content
                            full_content += content_chunk
                            result_queue.put(f"data: {json.dumps({'content': content_chunk})}\n\n")
                    result_queue.put('__DONE__')
                    if session_id:
                        result_queue.put(f"data: {json.dumps({'session_id': session_id})}\n\n")
                        try:
                            session = ChatSession.objects.get(session_id=session_id, user=request.user)
                            session.messages = messages + [{'role': 'assistant', 'content': full_content}]
                            session.save()
                        except ChatSession.DoesNotExist:
                            pass
                except Exception as e:
                    result_queue.put(f"data: {json.dumps({'error': str(e)})}\n\n")
                    result_queue.put('__DONE__')

            thread = threading.Thread(target=read_stream)
            thread.daemon = True
            thread.start()
            
            def event_stream():
                while True:
                    item = result_queue.get()
                    if item == '__DONE__':
                        break
                    yield item
            
            return StreamingHttpResponse(event_stream(), content_type='text/event-stream')
        else:
            result = {
                'content': response.choices[0].message.content,
                'model': response.model
            }
            
            if session_id and request.user.is_authenticated:
                result['session_id'] = session_id
                try:
                    session = ChatSession.objects.get(session_id=session_id, user=request.user)
                    session.messages = messages + [{'role': 'assistant', 'content': result['content']}]
                    session.save()
                except ChatSession.DoesNotExist:
                    pass
            
            return JsonResponse(result)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(['GET'])
def get_models(request):
    models = get_available_models()
    print(f"get_models() returning {len(models)} models: {models}")
    return JsonResponse({'models': models})


@csrf_exempt
@require_http_methods(['GET'])
def get_user_sessions(request):
    from .models import ChatSession
    if request.user.is_authenticated:
        sessions = ChatSession.objects.filter(user=request.user).order_by('-updated_at')[:20]
    else:
        sessions = ChatSession.objects.filter(user__isnull=True).order_by('-updated_at')[:20]
    sessions_list = [
        {
            'session_id': session.session_id,
            'title': session.title,
            'created_at': session.created_at.isoformat(),
            'updated_at': session.updated_at.isoformat()
        }
        for session in sessions
    ]
    return JsonResponse({'sessions': sessions_list})


@csrf_exempt
@require_http_methods(['POST'])
def create_session(request):
    if request.user.is_authenticated:
        session_id = str(uuid.uuid4())
        from .models import ChatSession
        session = ChatSession.objects.create(
            session_id=session_id,
            user=request.user,
            messages=[]
        )
        return JsonResponse({'session_id': session_id})
    else:
        return JsonResponse({'error': 'Unauthorized'}, status=401)


@csrf_exempt
@require_http_methods(['POST'])
def delete_session(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)
    
    try:
        data = json.loads(request.body)
        session_id = data.get('session_id')
        
        from .models import ChatSession
        ChatSession.objects.filter(session_id=session_id, user=request.user).delete()
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


import os
import re
import mimetypes

PROJECT_ROOT = str(BASE_DIR)
SKIP_DIRS = {'.git', '__pycache__', 'node_modules', '.venv', 'venv', 'env', 'build', 'dist', '.cache', '.mypy_cache'}
SKIP_EXTS = {'.pyc', '.pyo', '.o', '.so', '.dll', '.dylib', '.a', '.lib', '.exe', '.bin', '.db', '.sqlite', '.sqlite3', '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.svg', '.mp3', '.mp4', '.avi', '.wav', '.zip', '.tar', '.gz', '.7z', '.rar'}

def is_binary_file(filepath):
    try:
        with open(filepath, 'rb') as f:
            chunk = f.read(8192)
            if b'\x00' in chunk:
                return True
        return False
    except Exception:
        return True

def list_dir(path, base=None):
    if base is None:
        base = PROJECT_ROOT
    result = []
    try:
        entries = sorted(os.listdir(path))
    except PermissionError:
        return result
    except FileNotFoundError:
        return result
    for entry in entries:
        full_path = os.path.join(path, entry)
        rel_path = os.path.relpath(full_path, base)
        if os.path.isdir(full_path):
            if entry in SKIP_DIRS or entry.startswith('.'):
                continue
            sub_result = list_dir(full_path, base)
            if sub_result:
                result.append({'name': entry, 'type': 'dir', 'path': rel_path, 'children': sub_result})
            else:
                result.append({'name': entry, 'type': 'dir', 'path': rel_path, 'children': []})
        else:
            _, ext = os.path.splitext(entry)
            if ext.lower() in SKIP_EXTS:
                continue
            if is_binary_file(full_path):
                continue
            result.append({'name': entry, 'type': 'file', 'path': rel_path})
    return result


@csrf_exempt
@require_http_methods(['GET'])
def list_files_api(request):
    try:
        dir_path = request.GET.get('dir', '')
        if dir_path:
            full_path = os.path.join(PROJECT_ROOT, dir_path)
        else:
            full_path = PROJECT_ROOT
        files = list_dir(full_path)
        return JsonResponse({'files': files})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['POST'])
def set_workspace_api(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)
    try:
        data = json.loads(request.body)
        path = data.get('path', '').strip()
        if not path:
            return JsonResponse({'error': 'Путь не указан'}, status=400)
        ok, result = save_workspace_path(path)
        if ok:
            global PROJECT_ROOT
            PROJECT_ROOT = result
            return JsonResponse({'success': True, 'path': result})
        return JsonResponse({'error': result}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['GET'])
def get_workspace_api(request):
    return JsonResponse({'path': PROJECT_ROOT})


@csrf_exempt
@require_http_methods(['POST'])
def upload_file_api(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)

    try:
        uploaded_files = request.FILES.getlist('files[]') or request.FILES.getlist('files')
        if not uploaded_files:
            return JsonResponse({'error': 'No files uploaded'}, status=400)

        uploaded_paths = []
        for uploaded_file in uploaded_files:
            dest_path = os.path.join(PROJECT_ROOT, str(uploaded_file))
            dest_dir = os.path.dirname(dest_path)
            if not os.path.isdir(dest_dir):
                os.makedirs(dest_dir, exist_ok=True)
            with open(dest_path, 'wb+') as dest:
                for chunk in uploaded_file.chunks():
                    dest.write(chunk)
            uploaded_paths.append(str(uploaded_file))

        return JsonResponse({'success': True, 'files': uploaded_paths})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['GET'])
def read_file_api(request):
    try:
        file_path = request.GET.get('path', '')
        if not file_path:
            return JsonResponse({'error': 'No path specified'}, status=400)
        full_path = os.path.join(PROJECT_ROOT, file_path)
        if not os.path.isfile(full_path):
            return JsonResponse({'error': 'File not found'}, status=404)
        _, ext = os.path.splitext(file_path)
        if is_binary_file(full_path):
            return JsonResponse({'error': 'Binary file'}, status=400)
        with open(full_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        lines = content.split('\n')
        line_numbers = list(range(1, len(lines) + 1))
        return JsonResponse({
            'path': file_path,
            'content': content,
            'lines': line_numbers,
            'total_lines': len(lines)
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['GET'])
def search_files_api(request):
    try:
        query = request.GET.get('q', '').strip().lower()
        if not query or len(query) < 2:
            return JsonResponse({'results': []})
        
        results = []
        for root, dirs, files in os.walk(PROJECT_ROOT):
            dirs[:] = [d for d in dirs if d not in SKIP_DIRS and not d.startswith('.')]
            for filename in files:
                _, ext = os.path.splitext(filename)
                if ext.lower() in SKIP_EXTS:
                    continue
                if is_binary_file(os.path.join(root, filename)):
                    continue
                rel_path = os.path.relpath(root, PROJECT_ROOT)
                try:
                    with open(os.path.join(root, filename), 'r', encoding='utf-8', errors='replace') as f:
                        content = f.read()
                    lines = content.split('\n')
                    for i, line in enumerate(lines):
                        if query in line.lower():
                            results.append({
                                'file': os.path.join(rel_path, filename) if rel_path != '.' else filename,
                                'line_num': i + 1,
                                'line': line.strip()[:200],
                                'context': ''
                            })
                            if len(results) >= 100:
                                return JsonResponse({'results': results})
                except Exception:
                    continue
        return JsonResponse({'results': results[:100]})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['POST'])
def load_directory_context(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)

    try:
        data = json.loads(request.body)
        dir_path = data.get('path', '').strip()
        if not dir_path:
            return JsonResponse({'error': 'No path specified'}, status=400)

        full_path = os.path.join(PROJECT_ROOT, dir_path)
        if not os.path.isdir(full_path):
            return JsonResponse({'error': 'Directory not found'}, status=404)

        if not full_path.startswith(str(PROJECT_ROOT)):
            return JsonResponse({'error': 'Invalid path'}, status=400)

        max_files = int(data.get('max_files', 50))
        max_total_size_mb = int(data.get('max_total_size_mb', 10))

        files = []
        total_size = 0

        for root, dirs, filenames in os.walk(full_path):
            dirs[:] = [d for d in dirs if d not in SKIP_DIRS and not d.startswith('.')]

            for filename in sorted(filenames):
                if len(files) >= max_files:
                    break

                _, ext = os.path.splitext(filename)
                if ext.lower() in SKIP_EXTS:
                    continue

                file_full_path = os.path.join(root, filename)
                try:
                    file_size = os.path.getsize(file_full_path)
                    if total_size + file_size > max_total_size_mb * 1024 * 1024:
                        continue

                    if is_binary_file(file_full_path):
                        continue

                    with open(file_full_path, 'r', encoding='utf-8', errors='replace') as f:
                        content = f.read()

                    if not content.strip():
                        continue

                    rel_path = os.path.relpath(file_full_path, PROJECT_ROOT)
                    files.append({
                        'path': rel_path,
                        'content': content,
                        'size': file_size,
                        'lines': len(content.split('\n'))
                    })
                    total_size += file_size

                except Exception:
                    continue

        return JsonResponse({
            'success': True,
            'directory': dir_path,
            'file_count': len(files),
            'total_size_bytes': total_size,
            'files': files
        })

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['POST'])
def load_session_messages(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)

    try:
        data = json.loads(request.body)
        session_id = data.get('session_id')

        from .models import ChatSession
        session = ChatSession.objects.get(session_id=session_id, user=request.user)
        return JsonResponse({
            'session_id': session_id,
            'messages': session.messages
        })
    except ChatSession.DoesNotExist:
        return JsonResponse({'error': 'Session not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['POST'])
def run_command_api(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)

    try:
        data = json.loads(request.body)
        command = data.get('command', '').strip()
        if not command:
            return JsonResponse({'error': 'No command provided'}, status=400)

        if ';' in command or '&&' in command or '||' in command or '`' in command:
            return JsonResponse({'error': 'Dangerous characters not allowed'}, status=400)

        allowed_prefixes = ['python3 ', 'python ', 'npm ', 'pip ', 'pip3 ', 'node ', 'npx ',
                           'ls ', 'cat ', 'head ', 'tail ', 'grep ', 'find ', 'echo ',
                           'rm ', 'mkdir ', 'cp ', 'mv ', 'git ', 'pytest ', 'django-admin',
                           'chmod ', 'du ', 'df ', 'ps ', 'docker ', 'make ']

        if not any(command.startswith(p) for p in allowed_prefixes):
            return JsonResponse({'error': 'Command not allowed'}, status=403)

        import subprocess
        result = subprocess.run(
            command,
            shell=True,
            capture_output=True,
            text=True,
            timeout=60,
            cwd=str(PROJECT_ROOT)
        )
        output = result.stdout[:50000]
        if result.returncode != 0:
            error = result.stderr[:10000]
            return JsonResponse({
                'output': output,
                'error': error,
                'returncode': result.returncode
            })
        return JsonResponse({
            'output': output,
            'returncode': 0
        })
    except subprocess.TimeoutExpired:
        return JsonResponse({'error': 'Command timed out (60s)'}, status=408)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['POST'])
def create_file_api(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)

    try:
        data = json.loads(request.body)
        file_path = data.get('path', '').strip()
        content = data.get('content', '')

        if not file_path:
            return JsonResponse({'error': 'No path provided'}, status=400)

        full_path = os.path.join(PROJECT_ROOT, file_path)
        if not full_path.startswith(str(PROJECT_ROOT)):
            return JsonResponse({'error': 'Invalid path'}, status=400)

        dir_name = os.path.dirname(full_path)
        if not os.path.isdir(dir_name):
            os.makedirs(dir_name, exist_ok=True)

        with open(full_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return JsonResponse({'success': True, 'path': file_path})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['POST'])
def update_file_api(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)

    try:
        data = json.loads(request.body)
        file_path = data.get('path', '').strip()
        content = data.get('content', '')

        if not file_path:
            return JsonResponse({'error': 'No path provided'}, status=400)

        full_path = os.path.join(PROJECT_ROOT, file_path)
        if not full_path.startswith(str(PROJECT_ROOT)):
            return JsonResponse({'error': 'Invalid path'}, status=400)

        if not os.path.isfile(full_path):
            return JsonResponse({'error': 'File not found'}, status=404)

        with open(full_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return JsonResponse({'success': True, 'path': file_path})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['POST'])
def generate_docx_api(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)

    try:
        data = json.loads(request.body)
        title = data.get('title', 'Документ')
        content = data.get('content', '')
        filename = data.get('filename', 'document.docx')

        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph()

        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            if para_text.startswith('# '):
                doc.add_heading(para_text[2:].strip(), level=2)
            elif para_text.startswith('## '):
                doc.add_heading(para_text[3:].strip(), level=3)
            elif para_text.startswith('- '):
                lines = para_text.split('\n')
                for line in lines:
                    if line.startswith('- '):
                        doc.add_paragraph(line[2:].strip(), style='List Bullet')
            else:
                doc.add_paragraph(para_text.strip())

        import io
        from django.http import HttpResponse

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        safe_filename = ''.join(c for c in filename if c.isalnum() or c in ' _-').rstrip()
        if not safe_filename.endswith('.docx'):
            safe_filename += '.docx'

        response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{safe_filename}"'
        return response
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['POST'])
def save_session(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)
    
    try:
        data = json.loads(request.body)
        session_id = data.get('session_id')
        messages = data.get('messages', [])
        
        from .models import ChatSession
        session, created = ChatSession.objects.get_or_create(
            session_id=session_id,
            user=request.user
        )
        session.messages = messages
        if not session.title:
            user_messages = [m for m in messages if m.get('role') == 'user']
            if user_messages:
                first_msg = user_messages[0].get('content', '')[:60]
                session.title = first_msg.replace('\n', ' ')
            else:
                session.title = 'Новая сессия'
        session.save()
        
        return JsonResponse({'success': True, 'session_id': session_id})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['POST'])
def apply_patch_api(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)

    try:
        data = json.loads(request.body)
        file_path = data.get('path', '').strip()
        new_content = data.get('content', '')

        if not file_path:
            return JsonResponse({'error': 'No path provided'}, status=400)

        full_path = os.path.join(PROJECT_ROOT, file_path)
        if not full_path.startswith(str(PROJECT_ROOT)):
            return JsonResponse({'error': 'Invalid path'}, status=400)

        dir_name = os.path.dirname(full_path)
        if not os.path.isdir(dir_name):
            os.makedirs(dir_name, exist_ok=True)

        old_content = ''
        if os.path.isfile(full_path):
            with open(full_path, 'r', encoding='utf-8') as f:
                old_content = f.read()

        with open(full_path, 'w', encoding='utf-8') as f:
            f.write(new_content)

        import difflib
        old_lines = old_content.splitlines(keepends=True)
        new_lines = new_content.splitlines(keepends=True)
        diff = difflib.unified_diff(
            old_lines, new_lines,
            fromfile='a/' + file_path,
            tofile='b/' + file_path
        )
        diff_text = ''.join(diff)

        return JsonResponse({
            'success': True,
            'path': file_path,
            'diff': diff_text
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(['POST'])
def diff_api(request):
    if not request.user.is_authenticated:
        return HttpResponseForbidden({'error': 'Unauthorized'}, status=401)

    try:
        data = json.loads(request.body)
        old_content = data.get('old_content', '')
        new_content = data.get('new_content', '')
        filename = data.get('filename', 'file')

        import difflib
        old_lines = old_content.splitlines(keepends=True)
        new_lines = new_content.splitlines(keepends=True)
        diff = difflib.unified_diff(
            old_lines, new_lines,
            fromfile='a/' + filename,
            tofile='b/' + filename
        )
        diff_text = ''.join(diff)

        diff_lines = diff_text.splitlines()
        colorized = []
        for line in diff_lines:
            if line.startswith('---') or line.startswith('+++'):
                colorized.append({'type': 'header', 'text': line})
            elif line.startswith('@@'):
                colorized.append({'type': 'hunk', 'text': line})
            elif line.startswith('+'):
                colorized.append({'type': 'added', 'text': line[1:]})
            elif line.startswith('-'):
                colorized.append({'type': 'removed', 'text': line[1:]})
            else:
                colorized.append({'type': 'context', 'text': line})

        return JsonResponse({
            'diff': diff_text,
            'colorized': colorized
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# =====================================================================
# RAG / Vector Database API
# =====================================================================

from .services.rag_service import RagService

# Ленивая инициализация (чтобы не ломать Django без LM Studio)
_rag_service = None

def _get_rag_service():
    global _rag_service
    if _rag_service is None:
        try:
            _rag_service = RagService()
        except Exception:
            pass
    return _rag_service


@csrf_exempt
@require_http_methods(["GET"])
@login_required
def vector_db_list(request):
    """Возвращает список всех векторных баз данных"""
    rag = _get_rag_service()
    if rag is None:
        return JsonResponse({'error': 'RAG сервис не доступен (проверьте LM Studio)'}, status=503)

    try:
        databases = rag.get_ready_databases()
        return JsonResponse({'databases': databases})
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"vector_db_list error: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
@login_required
def vector_db_create(request):
    """Создаёт новую векторную базу данных"""
    rag = _get_rag_service()
    if rag is None:
        return JsonResponse({'error': 'RAG сервис не доступен'}, status=503)

    try:
        body = json.loads(request.body)
        db_name = body.get('name', '').strip()

        if not db_name:
            return JsonResponse({'error': 'Укажите имя базы данных'}, status=400)

        # Проверка на недопустимые символы
        if not re.match(r'^[a-zA-Z0-9_\-]+$', db_name):
            return JsonResponse({'error': 'Имя может содержать только буквы, цифры, _ и -'}, status=400)

        rag.vector_db_service.create_db(db_name)
        return JsonResponse({'success': True, 'name': db_name})
    except ValueError as e:
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"vector_db_create error: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
@login_required
def vector_db_delete(request):
    """Удаляет векторную базу данных"""
    rag = _get_rag_service()
    if rag is None:
        return JsonResponse({'error': 'RAG сервис не доступен'}, status=503)

    try:
        body = json.loads(request.body)
        db_name = body.get('name', '').strip()

        if not db_name:
            return JsonResponse({'error': 'Укажите имя базы данных'}, status=400)

        rag.vector_db_service.delete_db(db_name)
        return JsonResponse({'success': True})
    except ValueError as e:
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"vector_db_delete error: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
@login_required
def vector_db_info(request):
    """Возвращает информацию о векторной базе"""
    rag = _get_rag_service()
    if rag is None:
        return JsonResponse({'error': 'RAG сервис не доступен'}, status=503)

    try:
        body = json.loads(request.body)
        db_name = body.get('name', '').strip()

        if not db_name:
            return JsonResponse({'error': 'Укажите имя базы данных'}, status=400)

        info = rag.vector_db_service.get_db_info(db_name)
        return JsonResponse({'info': info})
    except ValueError as e:
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"vector_db_info error: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["GET"])
@login_required
def vector_db_files(request):
    """Возвращает список файлов для индексации"""
    rag = _get_rag_service()
    if rag is None:
        return JsonResponse({'error': 'RAG сервис не доступен'}, status=503)

    try:
        db_name = request.GET.get('name', '').strip()

        if not db_name:
            return JsonResponse({'error': 'Укажите имя базы данных'}, status=400)

        files = rag.vector_db_service.list_files_for_indexing(db_name)
        return JsonResponse({'files': files})
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"vector_db_files error: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
@login_required
def vector_db_upload_file(request):
    """
    Загружает файл в директорию Files векторной базы.
    Файл можно отправить через FormData или бинарные данные.
    """
    rag = _get_rag_service()
    if rag is None:
        return JsonResponse({'error': 'RAG сервис не доступен'}, status=503)

    try:
        db_name = request.POST.get('db_name', '').strip()
        file = request.FILES.get('file')

        if not db_name or not file:
            return JsonResponse({'error': 'Укажите базу данных и файл'}, status=400)

        # Проверка расширения
        supported = rag.document_parser.get_supported_extensions()
        ext = Path(file.name).suffix.lower()
        if ext not in supported:
            return JsonResponse({
                'error': f'Неподдерживаемый формат: {ext}. Поддерживаются: {", ".join(supported)}'
            }, status=400)

        # Сохраняем файл в Files/
        db_path = rag.vector_db_service.get_db_path(db_name)
        files_dir = db_path / 'Files'
        files_dir.mkdir(parents=True, exist_ok=True)

        file_path = files_dir / file.name

        # Если файл с таким именем уже есть, добавляем префикс
        if file_path.exists():
            file_path = files_dir / f"{uuid.uuid4().hex[:8]}_{file.name}"

        # Записываем файл
        with open(file_path, 'wb') as f:
            for chunk in file.chunks():
                f.write(chunk)

        # Добавляем в метаданные
        relative_path = str(file_path.relative_to(db_path))
        doc = rag.vector_db_service.add_document(db_name, str(file_path), relative_path)

        return JsonResponse({
            'success': True,
            'file_path': relative_path,
            'doc': doc
        })
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"vector_db_upload_file error: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
@login_required
def vector_db_index(request):
    """
    Индексирует файлы в векторной базе.
    Можно указать конкретные файлы или индексирует все.
    """
    rag = _get_rag_service()
    if rag is None:
        return JsonResponse({'error': 'RAG сервис не доступен'}, status=503)

    try:
        body = json.loads(request.body)
        db_name = body.get('name', '').strip()
        file_paths = body.get('file_paths', [])

        if not db_name:
            return JsonResponse({'error': 'Укажите имя базы данных'}, status=400)

        if file_paths:
            # Индексируем только указанные файлы
            for fp in file_paths:
                full_path = rag.vector_db_service.get_db_path(db_name) / 'Files' / fp
                if not full_path.exists():
                    return JsonResponse({'error': f'Файл не найден: {fp}'}, status=400)
                result = rag.index_file(db_name, str(full_path))
                if result['status'] == 'error':
                    return JsonResponse({'error': f"Ошибка: {result.get('error')}"}, status=400)

            return JsonResponse({'success': True, 'message': f'Индексировано {len(file_paths)} файлов'})
        else:
            # Индексируем все файлы
            result = rag.build_index(db_name)
            return JsonResponse({'success': True, **result})
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"vector_db_index error: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
@login_required
def vector_db_rebuild_index(request):
    """Полная перестройка индекса"""
    rag = _get_rag_service()
    if rag is None:
        return JsonResponse({'error': 'RAG сервис не доступен'}, status=503)

    try:
        body = json.loads(request.body)
        db_name = body.get('name', '').strip()

        if not db_name:
            return JsonResponse({'error': 'Укажите имя базы данных'}, status=400)

        result = rag.rebuild_index(db_name)
        return JsonResponse({'success': True, **result})
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"vector_db_rebuild_index error: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
@login_required
def vector_db_search(request):
    """
    Поиск в векторной базе.
    Возвращает релевантные чанки.
    """
    rag = _get_rag_service()
    if rag is None:
        return JsonResponse({'error': 'RAG сервис не доступен'}, status=503)

    try:
        body = json.loads(request.body)
        db_name = body.get('name', '').strip()
        query = body.get('query', '').strip()
        top_k = int(body.get('top_k', 5))

        if not db_name or not query:
            return JsonResponse({'error': 'Укажите базу данных и запрос'}, status=400)

        if top_k < 1 or top_k > 100:
            return JsonResponse({'error': 'top_k должен быть от 1 до 100'}, status=400)

        results = rag.search(db_name, query, top_k=top_k)
        return JsonResponse({'results': results})
    except ValueError as e:
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger = logging.getLogger(__name__)
        logger.error(f"vector_db_search error: {e}", exc_info=True)
        return JsonResponse({'error': str(e)}, status=500)
