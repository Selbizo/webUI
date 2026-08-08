import os
import json
import csv
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Попытка импорта_optional dependencies
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


class DocumentParser:
    """Парсер документов различных форматов в текст"""

    SUPPORTED_EXTENSIONS = {
        '.txt': '_parse_text',
        '.md': '_parse_text',
        '.markdown': '_parse_text',
        '.json': '_parse_json',
        '.csv': '_parse_csv',
        '.docx': '_parse_docx',
        '.pdf': '_parse_pdf',
        '.xlsx': '_parse_xlsx',
        '.xls': '_parse_xlsx',
    }

    def parse_file(self, file_path):
        """Парсит файл и возвращает текст"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Неподдерживаемый формат: {ext}")

        method = getattr(self, self.SUPPORTED_EXTENSIONS[ext])
        return method(file_path)

    def _parse_text(self, file_path):
        """Парсит текстовые файлы (txt, md)"""
        encodings = ['utf-8', 'cp1251', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Не удалось прочитать файл {file_path}: все кодировки исчерпаны")

    def _parse_json(self, file_path):
        """Парсит JSON файлы в текстовый вид"""
        content = self._parse_text(file_path)
        try:
            data = json.loads(content)
            return json.dumps(data, ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            return content

    def _parse_csv(self, file_path):
        """Парсит CSV файлы в текстовый вид"""
        encodings = ['utf-8', 'cp1251', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as f:
                    content = f.read()

                # Попытка прочитать как CSV
                reader = csv.reader(content.splitlines())
                rows = list(reader)

                if not rows:
                    return ""

                # Если больше одной колонки — возвращаем как таблицу
                if len(rows[0]) > 1:
                    result = []
                    for row in rows:
                        result.append(' | '.join(cell.strip() for cell in row))
                    return '\n'.join(result)
                else:
                    return content
            except Exception:
                continue
        return self._parse_text(file_path)

    def _parse_docx(self, file_path):
        """Парсит Word документы (.docx)"""
        if not HAS_DOCX:
            logger.warning("python-docx не установлен. Установите: pip install python-docx")
            raise ImportError("Установите python-docx: pip install python-docx")

        try:
            doc = DocxDocument(str(file_path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(' | '.join(cells))

            return '\n\n'.join(paragraphs)
        except Exception as e:
            raise ValueError(f"Ошибка чтения DOCX: {e}")

    def _parse_pdf(self, file_path):
        """Парсит PDF файлы через PyMuPDF"""
        if not HAS_PYMUPDF:
            logger.warning("PyMuPDF не установлен. Установите: pip install PyMuPDF")
            raise ImportError("Установите PyMuPDF: pip install PyMuPDF")

        try:
            doc = fitz.open(str(file_path))
            pages = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text.strip())
            doc.close()
            return '\n\n'.join(pages)
        except Exception as e:
            raise ValueError(f"Ошибка чтения PDF: {e}")

    def _parse_xlsx(self, file_path):
        """Парсит Excel файлы (.xlsx, .xls)"""
        if not HAS_OPENPYXL:
            logger.warning("openpyxl не установлен. Установите: pip install openpyxl")
            raise ImportError("Установите openpyxl: pip install openpyxl")

        try:
            wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
            sheets = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_lines = []

                for row in sheet.iter_rows(values_only=True):
                    cells = [str(cell) if cell is not None else '' for cell in row]
                    # Фильтруем пустые строки
                    if any(c.strip() for c in cells):
                        sheet_lines.append(' | '.join(c.strip() for c in cells))

                if sheet_lines:
                    sheets.append(f"## Лист: {sheet_name}\n" + '\n'.join(sheet_lines))

            wb.close()
            return '\n\n'.join(sheets)
        except Exception as e:
            raise ValueError(f"Ошибка чтения XLSX: {e}")

    def get_supported_extensions(self):
        """Возвращает список поддерживаемых расширений"""
        return list(self.SUPPORTED_EXTENSIONS.keys())

    def list_document_files(self, directory):
        """
        Список всех документов в директории (рекурсивно).
        Возвращает список кортежей (path, name, size, ext).
        """
        files = []
        directory = Path(directory)

        for ext, _ in self.SUPPORTED_EXTENSIONS.items():
            for file_path in directory.rglob(f'*{ext}'):
                if file_path.is_file():
                    files.append({
                        'path': str(file_path),
                        'name': file_path.name,
                        'relative_path': str(file_path.relative_to(directory)),
                        'size': file_path.stat().st_size,
                        'extension': ext,
                    })

        # Сортируем по пути
        files.sort(key=lambda f: f['relative_path'])
        return files
